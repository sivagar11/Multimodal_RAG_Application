"""
Report Generation Service
Handles creating financial reports from chat responses
"""
import pandas as pd
from docx import Document
from docx2pdf import convert
from langchain.chat_models import ChatOpenAI
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate

from backend.core.config import Config
from backend.core.logger import setup_logger

logger = setup_logger(__name__)


PROMPT_TEMPLATE = """
You are financial analyst and summary report writer using the provided informations.
Answer the question based only on the following context, which is from annual financial report of a tea export company called ceylon tea brokers.
{context}
Question: {question}
write the report with a suitable heading and use bullet points and other report writing technics.
Answer:
"""


class ReportService:
    """Service for generating financial reports"""
    
    def __init__(self):
        """Initialize report service"""
        self.config = Config
        self.csv_file = self.config.REPORTS_PATH / "report_data.csv"
        self.docx_file = self.config.REPORTS_PATH / "economic_analysis_report.docx"
        self.pdf_file = self.config.REPORTS_PATH / "report.pdf"
    
    def merge_last_n_rows(self, n: int) -> str:
        """
        Merge last N rows from report data CSV
        
        Args:
            n: Number of last rows to merge
            
        Returns:
            Merged text content
        """
        logger.info(f"Merging last {n} rows from report data")
        
        df = pd.read_csv(self.csv_file)
        
        if "Response" not in df.columns:
            raise ValueError("The 'Response' column is not present in the CSV file.")
        
        last_n_rows = df["Response"].tail(n)
        context = last_n_rows.str.cat(sep=" ")
        
        logger.info("Rows merged successfully")
        return context
    
    def answer(self, context: str) -> str:
        """
        Generate report answer using LLM
        
        Args:
            context: Context text
            
        Returns:
            Generated report text
        """
        logger.info("Generating report with LLM")
        
        question = "summarize these texts and give me the organized report"
        
        qa_chain = LLMChain(
            llm=ChatOpenAI(
                model=self.config.DEFAULT_LLM_MODEL,
                openai_api_key=self.config.OPENAI_API_KEY,
                max_tokens=3000
            ),
            prompt=PromptTemplate.from_template(PROMPT_TEMPLATE),
        )
        
        result = qa_chain.run({"context": context, "question": question})
        return result
    
    def create_word_document(self, content: str):
        """
        Create Word document from content
        
        Args:
            content: Report content
        """
        logger.info("Creating Word document")
        
        doc = Document()
        sections = content.split("\n\n")
        
        # Add title
        if sections:
            title = sections[0]
            doc.add_heading(title, level=1)
        
        # Add content
        for section in sections[1:]:
            lines = section.split("\n")
            for line in lines:
                if line.startswith("- "):
                    doc.add_paragraph(line, style="ListBullet")
                else:
                    doc.add_paragraph(line)
            doc.add_paragraph()
        
        # Save document
        doc.save(str(self.docx_file))
        logger.info(f"Word document created: {self.docx_file}")
    
    def convert_docx_to_pdf(self):
        """Convert DOCX to PDF"""
        logger.info("Converting Word to PDF")
        convert(str(self.docx_file), str(self.pdf_file))
        logger.info(f"PDF created: {self.pdf_file}")
    
    def report_generation(self, n: int) -> bool:
        """
        Complete report generation workflow
        
        Args:
            n: Number of last responses to include
            
        Returns:
            True if successful
        """
        logger.info(f"Starting report generation with n={n}")
        
        try:
            context = self.merge_last_n_rows(n)
            result = self.answer(context)
            self.create_word_document(result)
            self.convert_docx_to_pdf()
            
            logger.info("Report generation completed successfully")
            return True
        except Exception as e:
            logger.error(f"Error generating report: {e}")
            raise


# Create singleton instance
report_service = ReportService()

